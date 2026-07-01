#!/usr/bin/env python3
"""
Generates a technical reference document for the GPS track filtering pipeline.
Source of truth: lib/features/home/utils/trim_track.dart (docstring + implementation)
and lib/features/home/utils/filter_settings.dart / compute_daily_stats.dart.
Run with: python3 generate_filtering_doc.py
Output:  GPS-Track-Filtering-Pipeline.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette (matches Logbook_Upgrade_Plan.docx) ─────────────────────────
NAVY  = RGBColor(0x1A, 0x35, 0x5E)
TEAL  = RGBColor(0x0B, 0x6E, 0x8A)
SLATE = RGBColor(0x44, 0x55, 0x6B)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = TEAL
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = SLATE
    return h


def add_para(doc, text, bold=False, italic=False, color=SLATE, size=10.5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text, level=0, color=SLATE):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p


def add_numbered(doc, text, color=SLATE):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1A355E')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)

    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = SLATE

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4D8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = SLATE

# ── Cover ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run('GPS Track Filtering Pipeline')
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('Technical Reference — Logbook App')
r2.font.size = Pt(15)
r2.font.color.rgb = TEAL

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(
    f"Prepared: {datetime.date.today().strftime('%B %Y')}   ·   "
    f"Source: lib/features/home/utils/trim_track.dart"
)
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

doc.add_page_break()

# ── 1. Purpose ────────────────────────────────────────────────────────────────
add_heading(doc, '1. Purpose', level=1)
add_para(doc,
    'A raw GPX track from a phone or chartplotter is noisy in ways that are specific '
    'to marine use: the GPS receiver takes time to acquire a solid fix after power-on '
    '("cold start"), a single wrong fix sometimes appears before the receiver settles, '
    'stationary periods at anchor or on a mooring drift within a small radius rather '
    'than sitting at one exact point, and occasional teleports or dropouts produce '
    'physically impossible speed spikes. Rendering the raw points directly on a map '
    'produces a track that visibly starts in the wrong place, jitters while moored, '
    'and occasionally jumps across the chart.')
add_para(doc,
    'The filtering pipeline in trim_track.dart cleans a raw List<TrackPoint> into '
    'a form suitable for map rendering and for trip statistics, without discarding '
    'the information needed to show the user where and for how long the boat stopped.')

add_heading(doc, '2. Pipeline order', level=1)
add_para(doc, 'Each run of the pipeline (trimTrackWithAnchors / buildDisplayModel) executes the following seven passes, in this order:')
stages = [
    ('1. Annotate', 'Initial windowed speed + positional spread per fix, using a simple centred window (no flagged fixes exist yet).'),
    ('2. Spike detection', 'Flag physically implausible moving fixes — done BEFORE stop detection so a spike cannot masquerade as, or corrupt, a stop.'),
    ('3. Bad-first-fix detection', 'Flag fix[0] if it is a single position outlier relative to where the receiver settled by fix[1] onward.'),
    ('4. Re-annotate', 'Repeat pass 1, but the speed/spread window now stops at any flagged-fix boundary (gap-aware), since all flagged positions are now known.'),
    ('5. Stationary-segment detection', 'Find every stationary run of fixes; spread calculations exclude flagged fixes so a bad first fix cannot inflate the apparent anchor-swing radius.'),
    ('6. Cold-start detection', 'Within the start-of-track stop only, flag the leading fixes that sit anomalously far from the settled berth position — the GPS warm-up period.'),
    ('7. Smoothing', 'Apply a sliding-median filter (lat/lon independently) to the kept moving fixes, removing residual GPS jitter without distorting the track shape.'),
]
make_table(doc, ['Pass', 'What it does'], stages, col_widths=[1.7, 4.6])

add_divider(doc)

# ── 3. Geometry & data structures ────────────────────────────────────────────
add_heading(doc, '3. Core data structures', level=1)
add_para(doc, 'Distances use the Haversine great-circle formula throughout (Earth radius 6,371,000 m). Each raw TrackPoint is wrapped internally in a mutable annotation record during processing:')
add_code(doc, 'lat, lon, time            — from the source GPX fix')
add_code(doc, 'instSpeedKn               — instantaneous speed from the previous fix')
add_code(doc, 'winSpeedKn, winSpreadM    — centred-window speed / positional spread')
add_code(doc, 'stationary   (bool)       — inside a validated stop segment')
add_code(doc, 'flagged      (bool)       — spike or bad-first-fix; excluded from the cleaned track')
add_code(doc, 'coldStart    (bool)       — GPS warm-up fix at the very start of the track')

add_heading(doc, '4. Spike detection (pass 2)', level=1)
add_para(doc,
    'Only fixes classified as "moving" (not already stationary) are considered. The '
    'median instantaneous speed and the median absolute deviation (MAD) of that speed '
    'are computed across all moving fixes. A fix is flagged as a spike if its '
    'instantaneous speed exceeds:')
add_code(doc, 'max(maxSpeedKn, median + 4.0 × 1.4826 × MAD)')
add_para(doc,
    'The 1.4826 factor converts MAD to a standard-deviation-equivalent for a normal '
    'distribution; 4.0 is the number of MAD-sigmas allowed above the boat’s own '
    'recent speed distribution before a fix is considered implausible. maxSpeedKn is a '
    'hard floor on top of that statistical threshold — a fix faster than this is always '
    'flagged, regardless of how fast the boat has otherwise been travelling. This value '
    'is user-configurable (FilterSettings.maxSpeedKn, default 12 kn, adjustable to 60 kn '
    'in Settings → Track Filter) because the correct ceiling depends heavily on the boat: '
    '12 kn comfortably covers a cruising sailboat, but a racing yacht, foiler, or '
    'powerboat can genuinely exceed it, and a fixed ceiling would false-flag real speed '
    'as a GPS glitch.')

add_heading(doc, '5. Bad-first-fix detection (pass 3)', level=1)
add_para(doc,
    'Detects a specific, observed failure mode distinct from cold-start drift: a single '
    'wrong position at fix[0] that snaps immediately to the correct position at fix[1], '
    'rather than converging gradually over several fixes. Method: compute the centroid '
    'and mean spread of the next 10 fixes (fix[1..10]); if fix[0] sits further than '
    '5× that spread AND more than 30 m from the centroid, it is flagged. The '
    'tight-cluster condition (5× spread) is what prevents this from ever firing on '
    'a trip that genuinely starts under way, where the lookahead cluster is naturally '
    'spread out along the track rather than clustered at a berth.')
add_para(doc,
    'This behaviour was tuned against real Idefix log data: 3 of 9 real-world logs '
    'exhibited this exact pattern, with the bad first fix at 256 m (28 Jun), 89 m '
    '(18 Apr), and 75 m (02 May) from the settled position.',
    italic=True)

add_heading(doc, '6. Stationary-segment detection (pass 5)', level=1)
add_para(doc, 'A fix is classified stationary according to FilterSettings.stationaryMode:')
add_bullet(doc, 'speed mode (default) — windowed speed below speedThresholdKn (0.5 kn). A boat swinging at anchor through a wide arc, but at near-zero net speed, still collapses to a single position.')
add_bullet(doc, 'both mode — windowed speed AND windowed positional spread must both be below their thresholds (0.5 kn and 6 m). A wide anchor swing is NOT collapsed to one marker — use for anchor-watch style display.')
add_para(doc,
    'Consecutive stationary runs separated by 4 or fewer moving fixes are merged '
    '(bridges brief GPS flicker in/out of the stationary threshold) — unless the gap '
    'between them contains a flagged spike, which is never bridged, since that gap is a '
    'genuine teleport or dropout, not flicker.')
add_para(doc, 'A merged run is only accepted as a validated stop if:')
add_numbered(doc, 'Its duration is at least minStopMinutes (default 5 minutes) — filters out brief tacking lulls.')
add_numbered(doc, 'Its maximum radius from centroid (excluding flagged fixes) is at most maxStopSpreadM (default 30 m) — filters out slow drifting-in-light-air that dips below the speed threshold while still covering hundreds of metres.')
add_para(doc, 'Accepted stops are classified as start, mid, or end based on their position in the track (a stop beginning at or before the first non-flagged fix is "start" even if fix[0] itself was flagged as a bad first fix).')

add_heading(doc, '7. Cold-start detection (pass 6)', level=1)
add_para(doc,
    'Runs only against the start stop, if one was detected. Takes the second half of '
    'that stop’s fixes as the "settled" cloud and computes its centroid, mean '
    'distance, and standard deviation. The threshold radius is:')
add_code(doc, 'settled_mean + coldStartSettleFactor × max(settled_stddev, 0.5 m)')
add_para(doc,
    'Walking forward from fix[0], each consecutive fix beyond that threshold is flagged '
    'coldStart, stopping at the first fix that falls inside it (capped at 20 leading '
    'fixes) — cold-start convergence is monotonic in practice, so a fix that has already '
    'settled is never re-flagged later. coldStartSettleFactor defaults to 3.0; lower '
    'values are more aggressive about excluding leading fixes. Cold-start fixes are '
    'already inside the start stop (so excluded from the cleaned moving track either '
    'way), but the separate flag lets the UI show a distinct "GPS warm-up" indicator and '
    'excludes them from the anchor-position/spread statistics so the displayed stop '
    'position reflects the settled fix, not the noisy approach to it.')

add_divider(doc)

# ── 8. Stop anchors ───────────────────────────────────────────────────────────
add_heading(doc, '8. Stop anchors (position + spread)', level=1)
add_para(doc, 'Each validated stop is summarised as a TrackAnchor for map rendering (a "GPS halo"):')
add_bullet(doc, 'Centroid — mean lat/lon of the cluster (cold-start fixes excluded when at least 3 settled fixes remain).')
add_bullet(doc, 'CEP50 — radius containing 50% of the fixes (inner ring), floored at 10 m so the marker is always visible.')
add_bullet(doc, 'R95 — radius containing 95% of the fixes (outer ring), floored at 20 m.')

add_heading(doc, '9. Display segmentation & GPS-uncertainty bands', level=1)
add_para(doc,
    'buildDisplayModel walks the annotated fixes and emits a sequence of TrackSegments '
    'for rendering: moving runs, short stopEntry/stopExit connectors fading into a stop '
    'marker, and teleportBreak sentinels (no polyline drawn) at any point where a spike '
    'sits near a stop boundary and the jump exceeds 200 m — the visual gap is the signal '
    'that data is missing there, rather than drawing a straight line across it.')
add_para(doc,
    'For the moving segments, a per-fix GPS uncertainty band can also be rendered: a '
    'receiver’s base accuracy is inferred from the median R95 across all stops '
    '(default 8 m if there are none), and each moving fix’s local jitter — its '
    'deviation from the midpoint of its immediate neighbours — is combined with that '
    'base accuracy in quadrature, capped at 3× the base accuracy so one wild fix '
    'cannot blow the band to an implausible width. This is a display aid, meaningful '
    'only at close (harbour-level) map zoom.')

add_heading(doc, '10. Derived trip statistics', level=1)
add_para(doc, 'compute_daily_stats.dart runs the same pipeline and derives, for one day’s track:')
add_bullet(doc, 'Distance (nm) — sum of Haversine distances between consecutive kept moving fixes. Stops contribute zero distance.')
add_bullet(doc, 'Moving / stationary duration — stationary duration is the sum of validated stop durations; moving duration is the remainder.')
add_bullet(doc, 'Average over ground — moving distance ÷ moving time. An honest mechanical average including slow stretches.')
add_bullet(doc, 'Average making-way speed — same, but counting only fixes at or above makingWayThresholdKn (default 1.0 kn) — "how fast when actually sailing," excluding drift.')
add_bullet(doc, 'Max speed — the topSpeedPercentile (default p99) of instantaneous moving speeds, to suppress a lone remaining GPS glitch; the true single-fix maximum is also reported alongside it so an unusually low p99-vs-max gap is visible if it matters.')

add_divider(doc)

# ── 11. Configuration reference ───────────────────────────────────────────────
add_heading(doc, '11. Configuration reference (FilterSettings)', level=1)
add_para(doc,
    'All thresholds are user-tunable in Settings → Track Filter, persisted via '
    'ThemeProvider, and synced to the active logbook’s Firestore document — so the '
    'tuning is shared by everyone on a logbook and follows one user between their own '
    'devices.')
settings_rows = [
    ('stationaryMode', 'speed', 'speed = windowed speed only; both = speed AND positional spread (anchor-watch mode)'),
    ('speedThresholdKn', '0.5 kn', 'Below this windowed speed, a fix is "not making way"'),
    ('spreadThresholdM', '6 m', 'Only consulted in "both" mode — windowed positional spread ceiling for "holding position"'),
    ('window', '5 fixes', 'Half-width of the centred window used for the speed/spread signals'),
    ('smoothWindow', '3 fixes', 'Sliding-median smoothing window for the kept moving track (1 = off)'),
    ('minStopMinutes', '5.0 min', 'Minimum duration for a stationary run to count as a real stop'),
    ('maxStopSpreadM', '30 m', 'Maximum cluster radius for a validated stop'),
    ('detectColdStart', 'true', 'Enable/disable GPS warm-up fix stripping at track start'),
    ('coldStartSettleFactor', '3.0', 'Std-dev multiplier defining the cold-start settle threshold'),
    ('detectBadFirstFix', 'true', 'Enable/disable single-outlier detection at fix[0]'),
    ('makingWayThresholdKn', '1.0 kn', 'Minimum speed counted toward the making-way average'),
    ('topSpeedPercentile', '0.99', 'Percentile of moving speeds reported as "max speed"'),
    ('maxSpeedKn', '12 kn (8–60 adjustable)', 'Hard speed ceiling — always flagged as a spike above this, regardless of the boat’s own speed distribution'),
]
make_table(doc, ['Setting', 'Default', 'Meaning'], settings_rows, col_widths=[1.6, 1.3, 3.4])

doc.save('GPS-Track-Filtering-Pipeline.docx')
print('Wrote GPS-Track-Filtering-Pipeline.docx')
