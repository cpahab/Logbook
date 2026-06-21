#!/usr/bin/env python3
"""
Generates the Logbook app upgrade planning document.
Run with: python3 generate_doc.py
Output:  Logbook_Upgrade_Plan.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x35, 0x5E)   # headings
TEAL      = RGBColor(0x0B, 0x6E, 0x8A)   # sub-headings
SLATE     = RGBColor(0x44, 0x55, 0x6B)   # body
LIGHT     = RGBColor(0xF0, 0xF4, 0xF8)   # table header bg (approximated via shading)
RED       = RGBColor(0xC0, 0x39, 0x2B)   # risk / warning
GREEN     = RGBColor(0x1A, 0x7A, 0x4A)   # recommendation / green
AMBER     = RGBColor(0xD4, 0x7A, 0x00)   # caution

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = TEAL
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = SLATE
    return h

def add_para(doc, text, bold=False, italic=False, color=SLATE, size=10.5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(doc, text, level=0, color=SLATE):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p

def add_numbered(doc, text, color=SLATE):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_callout(doc, label, text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}  ")
    r1.font.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.color.rgb = SLATE
    r2.font.size = Pt(10.5)
    r2.font.italic = True
    return p

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '1A355E')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F5F8FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = SLATE

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4D8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default paragraph font
doc.styles['Normal'].font.name  = 'Calibri'
doc.styles['Normal'].font.size  = Pt(10.5)
doc.styles['Normal'].font.color.rgb = SLATE

# ── Cover ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
r = title.add_run('Logbook App')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('Upgrade Planning & Architecture Document')
r2.font.size = Pt(16)
r2.font.color.rgb = TEAL

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %Y')}   ·   Branch: main   ·   Version target: v1.0.17+")
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

doc.add_page_break()

# ── Table of Contents (manual) ────────────────────────────────────────────────
add_heading(doc, 'Contents', level=1)
toc_items = [
    ('1', 'Current State of the App'),
    ('2', 'Completed Changes — v1.0.17'),
    ('3', 'Future Upgrade: Multilingual (de / en)'),
    ('4', 'Future Upgrade: Authentication'),
    ('5', 'Future Upgrade: Billing & Cost Strategy'),
    ('6', 'Maps — Required Change Before Any App Store Submission'),
    ('7', 'Cost Risk Analysis — Running the App for Free'),
    ('8', 'Recommended Roadmap & Sequencing'),
    ('A', 'Appendix A — Implementation Prompts for Future Sessions'),
    ('B', 'Appendix B — Firebase Cost Reference'),
    ('C', 'Appendix C — Key Architectural Decisions Log'),
]
for num, label in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {num}.  {label}")
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEAL

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CURRENT STATE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Current State of the App', level=1)

add_para(doc,
    'The Logbook app is a Flutter-based sailing logbook for iOS, macOS, and Android (Android partially configured). '
    'It is currently used only by the developer. The app stores data locally using Hive and optionally syncs '
    'to Firebase Firestore using an installation-code model — two devices share a logbook by exchanging an '
    'eight-character code. There is no user authentication.',
)

doc.add_paragraph()
add_heading(doc, 'Technology stack', level=2)
make_table(doc,
    ['Layer', 'Technology', 'Notes'],
    [
        ['Local storage', 'Hive (typed adapters)', 'DayEntry, TimelineEntry, CrewMember, DailyTrack, EmergencyContact'],
        ['Cloud sync', 'Firebase Firestore + Storage', 'Keyed by installationId — no auth required today'],
        ['State management', 'provider 6.1.2 (ChangeNotifier)', '3 providers: HomeRepository, ThemeProvider, EmergencyRepository'],
        ['Navigation', 'go_router 17.2.3', 'Declarative routes, GoRouter redirect hooks available'],
        ['Maps', 'flutter_map 8.0.0', 'OSM tiles + Esri satellite — both require tile provider change before App Store'],
        ['Auth package', 'firebase_auth (in pubspec)', 'Present but never wired up — zero auth code exists'],
        ['Platform targets', 'iOS ✓  macOS ✓  Android ⚠', 'Android Firebase not yet configured (placeholder appId)'],
        ['Localization', 'flutter_localizations + intl', 'Hardcoded to de_CH — no ARB files, no l10n.yaml'],
        ['Billing / IAP', 'None', 'Zero billing code exists'],
    ],
    col_widths=[1.4, 1.8, 3.4],
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — COMPLETED CHANGES v1.0.17
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Completed Changes — v1.0.17', level=1)
add_para(doc,
    'The following items were designed and implemented in the current development session and are '
    'committed to the main branch. They are documented here for reference and as a baseline for '
    'future sessions.')

doc.add_paragraph()
add_heading(doc, 'Crew Roster (Items 1–6)', level=2)
add_bullet(doc, 'Allergies and conditions fields added to CrewMember Hive model; persisted locally and synced to Firestore.')
add_bullet(doc, 'Empty crew list opens directly in add-member mode (no redundant edit step).')
add_bullet(doc, 'Non-empty crew list uses a pending-buffer pattern: edits are staged and committed only on "Übernehmen" — matching the vessel-status interaction.')
add_bullet(doc, 'First timeline entry of a voyage auto-captures the current crew list as a timestamped crew snapshot entry.')
add_bullet(doc, 'Any crew change after the first snapshot auto-creates a new timeline entry recording the delta.')
add_bullet(doc, 'Crew section uses ÄNDERN/HINZUFÜGEN labels matching the AKTUALISIEREN pattern on vessel status.')

doc.add_paragraph()
add_heading(doc, 'Timeline Entry (Items 7–9)', level=2)
add_bullet(doc, 'New log entries start blank — no field pre-fill from the previous entry.')
add_bullet(doc, 'Course field constrained to 0–359 via a custom TextInputFormatter; out-of-range values are rejected on keypress.')
add_bullet(doc, 'Deleting a log entry no longer clears unrelated snackbars (removed clearSnackBars() call).')

doc.add_paragraph()
add_heading(doc, 'Tooltips & GPS Consent (Items 10–11)', level=2)
add_bullet(doc, 'Tooltips added to all edit/add icon buttons: crew add/edit, Etappe erfassen, Etappe speichern, Logeintrag bearbeiten/löschen, Fotos hinzufügen, Schiffsstatus aktualisieren, and Emergency Manifest edit buttons.')
add_bullet(doc, 'GpsConsentService created: shows a purpose-explanation dialog at first app frame (HomeScreen.initState via addPostFrameCallback) before the OS permission prompt. Text is urgency-focused and explains the Mayday use case.')

doc.add_paragraph()
add_heading(doc, 'Emergency Manifest (Items 12–14)', level=2)
add_bullet(doc, 'Vessel safety info (MMSI, Rufzeichen, Rettungsinsel, EPIRB, Feuerlöschanlage) is now editable inline via an AlertDialog on the Emergency Manifest page — no longer navigates away to Settings. Settings editing path remains unchanged.')
add_bullet(doc, 'The redundant "Tap Edit to add safety equipment details" hint text was removed; the EDIT button in the section header is the single edit trigger.')
add_bullet(doc, 'Crew Medical Overview header now shows an info-icon tooltip explaining that crew data is auto-filled from the most recent logbook day entry.')

doc.add_paragraph()
add_heading(doc, 'Dashboard & Map Screen (Items 15–17)', level=2)
add_bullet(doc, 'Stat boxes harmonized: both Dashboard and Map screens now use Icons.sailing for Segeltage, identical typography (Newsreader 22pt, Inter 10pt label), border radius 12, subtle box shadow, and uppercase labels.')
add_bullet(doc, 'Map filter presets changed from JAHR/6 MON/3 MON/EIGENE to 1 JAHR/1 MONAT/1 WOCHE/EIGENE. Default is now 1 Monat.')
add_bullet(doc, 'Custom date chip uses compact d.M.yy format with reduced letter spacing; date range picker wrapped in MediaQuery.withClampedTextScaling to prevent font-scale clipping in picker input fields.')

add_divider(doc)
add_callout(doc, 'Branch:', 'All v1.0.17 changes are on main. The previous main is preserved as branch main-v1.0.16.')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — i18n
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Future Upgrade: Multilingual (de / en)', level=1)

add_heading(doc, 'Scope decision', level=2)
add_para(doc, 'German and English only. The Emergency Manifest screen stays in English permanently — it is already written in English and this is the correct call for a maritime safety document that may be used in international waters or by non-German-speaking crew.')

doc.add_paragraph()
add_heading(doc, 'Current state', level=2)
add_bullet(doc, 'flutter_localizations and intl packages are already in pubspec.yaml.')
add_bullet(doc, 'The app locale is hardcoded to de_CH in app.dart.')
add_bullet(doc, 'Approximately 80–90 German strings are scattered across presentation files. No ARB files or l10n.yaml exist.')
add_bullet(doc, 'One structural issue: the string "Besatzung: " serves dual purpose — it is displayed text AND a logic sentinel used in startsWith() detection. Localising it naively breaks crew-change detection.')

doc.add_paragraph()
add_heading(doc, 'The sentinel fix (must be done first)', level=2)
add_para(doc,
    'Before any string extraction, the "Besatzung: " prefix must be split into two things: '
    'a non-localised ASCII key used in code logic (e.g. "crew:") and a separate display string '
    'in the ARB file. The Firestore-stored vesselStatusNote values also use this prefix, '
    'so existing stored entries will need a one-time migration or the detection logic must '
    'handle both the old and new prefix.')

doc.add_paragraph()
add_heading(doc, 'Work phases', level=2)
make_table(doc,
    ['Phase', 'Work', 'Est. days'],
    [
        ['1.1 Infrastructure', 'Add l10n.yaml, configure gen-l10n, wire AppLocalizations delegate in app.dart, language toggle in Settings persisted to Hive', '2'],
        ['1.2 String extraction', 'Extract ~80–90 strings into app_de.arb, fix Besatzung sentinel, handle ICU interpolations (counts, distances)', '4'],
        ['1.3 English ARB', 'Create app_en.arb — many terms (nm, kn, MMSI, VHF, GPX) are already English; main work is UI copy', '2'],
        ['1.4 Date formatting', 'Verify DateFormat calls respect active locale; German: TT.MM.JJJJ, English: DD/MM/YYYY (maritime convention)', '1'],
        ['1.5 QA', 'All screens in both languages, PDF export, edge cases and empty states', '2'],
    ],
    col_widths=[1.5, 4.2, 0.9],
)
doc.add_paragraph()
add_callout(doc, 'Total:', '~11 working days. No dependencies — can start immediately.', color=GREEN)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — AUTH
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Future Upgrade: Authentication', level=1)

add_heading(doc, 'Decisions already made', level=2)
add_bullet(doc, 'Providers: Email/Password, Apple Sign-In, Google Sign-In.')
add_bullet(doc, 'No grace period for existing users (developer is the only current user — clean cutover).')
add_bullet(doc, 'Platforms: iOS, Android, macOS.')

doc.add_paragraph()
add_heading(doc, 'Data model change', level=2)
add_para(doc,
    'The current Firestore structure is keyed by an installationId with no access control — '
    'anyone who knows the code can read the logbook. Authentication enables proper access control.')

make_table(doc,
    ['', 'Today', 'After auth'],
    [
        ['Firestore path', 'logbooks/{installationId}/entries/…', 'boats/{boatId}/entries/…'],
        ['Access control', 'None — anyone with the code can read', 'Firestore rules: only members of the boat'],
        ['User identity', 'None', 'users/{uid}/profile → {boatId, email, createdAt}'],
        ['Multi-device share', 'Exchange 8-char installationId', 'Same boatId, second uid added to members array'],
        ['Migration needed?', 'N/A (developer only)', 'Clean start — no migration scripts required'],
    ],
    col_widths=[1.6, 2.5, 2.5],
)

doc.add_paragraph()
add_heading(doc, 'Android and macOS prerequisites', level=2)
add_para(doc, 'Before any auth work can be tested on Android or macOS:')
add_bullet(doc, 'Android: Run flutterfire configure --platforms=android, add google-services.json to android/app/, set applicationId in build.gradle.')
add_bullet(doc, 'macOS: Register a separate Firebase macOS app (currently shares the iOS appId — needs its own for correct Auth OAuth callbacks).')
add_bullet(doc, 'Sign in with Apple: add the capability to both iOS and macOS Xcode targets; add entitlement files.')
add_bullet(doc, 'Google Sign-In: add SHA-1 fingerprint to Firebase Console for Android; add reversed-client-id URL scheme to Info.plist for iOS.')

doc.add_paragraph()
add_heading(doc, 'Work phases', level=2)
make_table(doc,
    ['Phase', 'Work', 'Est. days'],
    [
        ['2.1 Platform setup', 'Firebase Android registration, macOS app separation, enable Auth providers, add sign_in_with_apple + google_sign_in packages, entitlements and URL schemes', '5'],
        ['2.2 Auth UI', 'Login screen (email + Apple + Google), registration, forgot-password, auth state listener driving GoRouter redirect, account screen in Settings (sign out, delete account)', '5'],
        ['2.3 Data model', 'New Firestore structure, Firestore Security Rules, FirestoreService refactor to use boatId from user profile', '5'],
        ['2.4 Share model', '"Connect another device" flow: enter boatId on second device → uid added to boats/{boatId}/members; disconnect / leave boat actions', '3'],
        ['2.5 Auth QA', 'All three providers on all three platforms, account linking edge case, offline auth state persistence', '2'],
    ],
    col_widths=[1.5, 4.2, 0.9],
)
doc.add_paragraph()
add_callout(doc, 'Total:', '~20 working days. Requires Phase 2.1 before any Android testing can begin.', color=AMBER)

doc.add_paragraph()
add_heading(doc, 'Account linking edge case', level=2)
add_para(doc,
    'If a user signs in with Google using the same email address as an existing email/password account, '
    'Firebase throws a credential-already-in-use error. The app must catch this and offer an '
    '"Link accounts" flow. This is a required UX path — it is surprisingly common and app reviewers test it.')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — BILLING
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Future Upgrade: Billing & Cost Strategy', level=1)

add_heading(doc, 'The core principle', level=2)
add_para(doc,
    'The developer does not want to subsidise cloud storage costs. However, the actual Firebase '
    'cost per user is very small — roughly €0.003–0.006 per active user per month. The question '
    'is therefore not "how much does it cost" but "what model avoids paywall complexity in the app '
    'while still recovering costs if the app grows."')

doc.add_paragraph()
add_heading(doc, 'Option comparison', level=2)
make_table(doc,
    ['Model', 'Dev cost', 'User cost', 'Paywall in app', 'Android', 'Effort'],
    [
        ['Absorb Firebase costs', '$0–20/mo', 'None', 'None', 'Yes', 'Zero'],
        ['CloudKit (Apple platforms)', '$0', 'None (iCloud quota)', 'None', 'No', '3 weeks'],
        ['Web subscription + token', '$0 (users pay on website)', '~€10/yr via Stripe', 'None — token entry only', 'Yes', '2 weeks'],
        ['Local network sync (Bonjour)', '$0', 'None', 'None', 'Yes', '3 weeks'],
        ['Full IAP paywall', 'Firebase costs', '~€2/yr via App Store', 'Yes — full Apple review', 'Yes', '6–8 weeks'],
    ],
    col_widths=[1.7, 1.1, 1.4, 1.4, 0.8, 0.9],
)

doc.add_paragraph()
add_heading(doc, 'Recommended approach: absorb costs now, add web subscription later if needed', level=2)
add_para(doc,
    'Given that the app currently has no users other than the developer, and given the negligible '
    'Firebase cost at early scale, the recommended strategy is:')
add_numbered(doc, 'Deploy without any billing. Set a Firebase budget alert at €20/month and a hard spending cap at €50/month. This protects against runaway bugs or viral moments.')
add_numbered(doc, 'At ~500 active users, evaluate whether the Firebase bill is meaningful. It will likely be €2–4/month.')
add_numbered(doc, 'If cost recovery is needed, implement the web subscription + token model: users pay on a webpage via Stripe (avoiding Apple\'s 30% cut and App Store IAP review), receive an email with a sync token, and enter it in Settings. This is explicitly permitted by Apple guidelines when the service is delivered outside the app.')

doc.add_paragraph()
add_callout(doc, 'Why not IAP now?',
    'App Store IAP adds 30% Apple cut, a dedicated review process for subscriptions, '
    'RevenueCat integration complexity, paywall UI, Cloud Function webhook infrastructure, '
    'and ongoing subscription management. The break-even point (where this complexity is '
    'worth the effort) is at several hundred paying users — not where the app is today.',
    color=AMBER)

doc.add_paragraph()
add_heading(doc, 'Web subscription + token model (detail)', level=2)
add_para(doc, 'If and when cost recovery is implemented, the web + token model works as follows:')
add_bullet(doc, 'User visits logbuch.app/sync in a browser and pays via Stripe (€9.90/year example price).')
add_bullet(doc, 'Stripe webhook calls a Firebase Cloud Function that generates a UUID sync token and emails it to the user.')
add_bullet(doc, 'User enters the token in the app Settings under "Sync-Code aktivieren."')
add_bullet(doc, 'The app sends the token to a Cloud Function that validates it and marks the user\'s Firestore subscription as active.')
add_bullet(doc, 'Firestore Security Rules gate write access on subscription.status == "active".')
add_para(doc, 'No in-app purchase UI, no Apple review for billing, no 30% platform fee.')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — MAPS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Maps — Required Change Before Any App Store Submission', level=1)

add_para(doc,
    'This section is distinct from auth and billing — it must be addressed regardless of which '
    'other upgrades are pursued, because both current tile providers violate their terms of service '
    'for published apps.')

doc.add_paragraph()
add_heading(doc, 'Current situation', level=2)
add_para(doc, 'The app uses two tile sources, called in five places (tracks screen ×2, day detail ×2, PDF exporter ×1):')
add_bullet(doc, 'OpenStreetMap: tile.openstreetmap.org — OSM\'s own servers are explicitly off-limits for distributed apps. Their policy states that heavy use (including app distribution) is forbidden without prior permission from the OSM Operations Working Group.')
add_bullet(doc, 'Esri satellite: server.arcgisonline.com — the ArcGIS Living Atlas World Imagery basemap requires an Esri Developer account and an API key for production use, even on the free tier. Currently called with no credentials.')

doc.add_paragraph()
add_heading(doc, 'Recommended solution: MapTiler', level=2)
add_para(doc,
    'MapTiler is an OSM-compatible tile CDN that is drop-in compatible with flutter_map. '
    'It is specifically recommended for this app because it offers a nautical chart map style '
    '(depth contours, buoys, traffic separation zones, port markers) — directly relevant to sailors. '
    'The free tier covers early-stage usage comfortably.')

make_table(doc,
    ['MapTiler tier', 'Tile requests/month', 'Cost'],
    [
        ['Free', '100,000', '€0'],
        ['Pay-as-you-go', 'Unlimited', '€0.003 per 1,000 above free tier'],
        ['At 500 users (est.)', '~50,000 requests', '€0 — within free tier'],
        ['At 2,000 users (est.)', '~200,000 requests', '~€0.30/month'],
    ],
    col_widths=[1.8, 2.2, 2.6],
)

doc.add_paragraph()
add_para(doc, 'For the Esri satellite layer, the existing URL continues to work — just add an API key from the free ArcGIS Developer tier (2 million tile requests/month free).')

doc.add_paragraph()
add_heading(doc, 'Code change', level=2)
add_para(doc, 'The change is a URL swap in five locations plus a constants file:')
add_code(doc, '// lib/core/constants/map_config.dart')
add_code(doc, "const kMapTilerKey = 'YOUR_MAPTILER_API_KEY';")
add_code(doc, "const kBaseTileUrl = 'https://api.maptiler.com/maps/nautical/{z}/{x}/{y}.png?key=\$kMapTilerKey';")
add_code(doc, "const kSatelliteUrl = 'https://server.arcgisonline.com/.../tile/{z}/{y}/{x}?token=\$kEsriKey';")

add_callout(doc, 'Effort:', '2 days. No architectural change — purely a URL and credentials swap.', color=GREEN)
add_callout(doc, 'Note:', 'The flutter_map package and latlong2 do not need to change. No provider migration required.', color=TEAL)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — COST RISK
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Cost Risk Analysis — Running the App for Free', level=1)

add_heading(doc, 'Firebase free tier (Spark plan) capacity', level=2)
make_table(doc,
    ['Firebase service', 'Free tier limit', 'App usage at 500 users', 'Headroom'],
    [
        ['Firestore reads', '50,000/day', '~5,000/day', '10× margin'],
        ['Firestore writes', '20,000/day', '~2,000/day', '10× margin'],
        ['Firestore storage', '1 GB total', '~365 MB (365 entries × 1 KB × 1,000 users)', 'Adequate'],
        ['Firebase Storage (GPX)', '5 GB total', '~5 GB (10 MB × 500 users)', 'At limit'],
        ['Firebase Auth (email)', '10,000/month', 'Trivial', 'Fine'],
        ['Auth (Google/Apple)', 'Unlimited on Identity Platform', '—', 'Free'],
    ],
    col_widths=[1.8, 1.6, 2.2, 1.2],
)

doc.add_paragraph()
add_heading(doc, 'Blaze (pay-as-you-go) cost estimates', level=2)
make_table(doc,
    ['Active users', 'Estimated monthly Firebase cost'],
    [
        ['100', '€0 (Spark free tier covers all)'],
        ['500', '€0–1'],
        ['1,000', '€3–6'],
        ['5,000', '€18–30'],
        ['10,000', '€50–80'],
    ],
    col_widths=[2.0, 4.6],
)

doc.add_paragraph()
add_heading(doc, 'Risk mitigation', level=2)
add_bullet(doc, 'Set a Firebase billing budget alert at €20/month (email notification).')
add_bullet(doc, 'Set a hard spending cap at €50/month (Firebase automatically disables billing — note: this also disables Firestore writes above the Spark limit, so users would see sync failures rather than unexpected charges).')
add_bullet(doc, 'A runaway read/write bug is the main risk. Firebase\'s billing console shows per-operation usage in near-real-time — monitor this for the first week after any sync-related code change.')

doc.add_paragraph()
add_callout(doc, 'Verdict:',
    'Running completely free is financially sound for the foreseeable life of a niche sailing logbook. '
    'Getting to 1,000 active users — where cost reaches €5/month — is itself a significant milestone. '
    'The entire first year of Firebase costs will likely be under €50 total.',
    color=GREEN)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ROADMAP
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Recommended Roadmap & Sequencing', level=1)

add_heading(doc, 'Immediate (before App Store submission)', level=2)
add_numbered(doc, 'Switch map tiles to MapTiler + Esri API key. (2 days — no dependencies)')
add_numbered(doc, 'Register Android app in Firebase Console; add google-services.json. (1 day — needed for multi-platform testing)')

doc.add_paragraph()
add_heading(doc, 'Short term (parallel tracks, can run concurrently)', level=2)
add_numbered(doc, 'i18n: de/en (11 days — independent of auth; Emergency Manifest stays English).')
add_numbered(doc, 'Auth: Firebase Auth with email, Apple, Google (20 days — begin Phase 2.1 concurrently with i18n).')

doc.add_paragraph()
add_heading(doc, 'Medium term (after auth)', level=2)
add_numbered(doc, 'App Store and Google Play submissions for all three platforms.')
add_numbered(doc, 'Set Firebase budget alerts. Deploy without billing.')

doc.add_paragraph()
add_heading(doc, 'Later (if/when needed)', level=2)
add_numbered(doc, 'Web subscription + token billing model — only if Firebase cost becomes real or cost recovery is desired.')
add_numbered(doc, 'Additional languages — once ARB infrastructure is in place, each new language is 2–3 days of translation.')
add_numbered(doc, 'Nautical chart enhancements — MapTiler nautical style + OpenSeaMap overlay for chart markers.')

doc.add_paragraph()
add_heading(doc, 'Full calendar view', level=2)
make_table(doc,
    ['Weeks', 'Work', 'Dependency'],
    [
        ['1', 'Maps tile switch + Android Firebase registration', 'None — start here'],
        ['1–3', 'i18n (parallel)', 'None'],
        ['2–3', 'Auth Phase 2.1 — platform setup', 'After Android registration'],
        ['4–5', 'Auth Phase 2.2–2.3 — auth UI + data model', 'After 2.1'],
        ['6', 'Auth Phase 2.4–2.5 — share model + QA', 'After 2.3'],
        ['7–8', 'App Store Connect + Google Play setup, submissions', 'After auth + i18n complete'],
        ['8+', 'Apple review clock (1–2 weeks)', 'After submission'],
        ['Later', 'Web billing (if needed)', 'Auth must be complete'],
    ],
    col_widths=[0.8, 3.8, 2.0],
)

doc.add_paragraph()
add_callout(doc, 'Total elapsed time to first public release:',
    '8–10 weeks (with i18n and auth Phase 2.1 running in parallel from week 1).',
    color=NAVY)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX A — PROMPTS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Appendix A — Implementation Prompts for Future Sessions', level=1)
add_para(doc,
    'Each prompt below is designed to be pasted directly into a new Claude Code session. '
    'They are self-contained — each includes the necessary context so the agent can begin '
    'work without reading back through previous sessions. '
    'The branch and file paths match the current repository state on main.')

doc.add_paragraph()
add_divider(doc)

# A1
add_heading(doc, 'A1 — Map Tile Provider Switch (MapTiler + Esri)', level=2)
add_para(doc, 'Effort: 2 days | Dependency: None | Do this first before any App Store submission.', italic=True, color=SLATE)
doc.add_paragraph()
add_para(doc, 'Paste this prompt into a new session:', bold=True)
add_code(doc, 'We are working on the Flutter sailing logbook app at /home/user/Logbook (Flutter,')
add_code(doc, 'Hive, Firebase, flutter_map). The app currently uses OSM demo tiles')
add_code(doc, '(tile.openstreetmap.org) and Esri satellite tiles (server.arcgisonline.com)')
add_code(doc, 'with no API keys. Both violate their terms for published apps.')
add_code(doc, '')
add_code(doc, 'Task: Switch to MapTiler (nautical chart style) + Esri (with API key).')
add_code(doc, '')
add_code(doc, '1. Create lib/core/constants/map_config.dart with:')
add_code(doc, '   - kMapTilerApiKey (placeholder string)')
add_code(doc, '   - kBaseTileUrl using MapTiler nautical style:')
add_code(doc, '     https://api.maptiler.com/maps/nautical/{z}/{x}/{y}.png?key={key}')
add_code(doc, '   - kSatelliteUrl using Esri World Imagery with API key param')
add_code(doc, '')
add_code(doc, '2. Replace the hardcoded URL strings in all 5 places:')
add_code(doc, '   - lib/features/tracks/presentation/tracks_screen.dart (×2)')
add_code(doc, '   - lib/features/home/presentation/day_detail_screen.dart (×2)')
add_code(doc, '   - lib/features/home/utils/pdf_exporter.dart (×1 static URL)')
add_code(doc, '')
add_code(doc, '3. Update RichAttributionWidget in both map widgets to credit MapTiler.')
add_code(doc, '')
add_code(doc, 'Do not change any other code. Commit and push to main.')

doc.add_paragraph()
add_divider(doc)

# A2
add_heading(doc, 'A2 — Android Firebase Registration', level=2)
add_para(doc, 'Effort: 1 day | Dependency: None | Required before any Android testing.', italic=True, color=SLATE)
doc.add_paragraph()
add_para(doc, 'Paste this prompt into a new session:', bold=True)
add_code(doc, 'We are working on the Flutter sailing logbook app at /home/user/Logbook.')
add_code(doc, 'Firebase is configured for iOS and macOS but Android has a placeholder')
add_code(doc, 'appId in lib/firebase_options.dart: "REPLACE_WITH_ANDROID_APP_ID".')
add_code(doc, '')
add_code(doc, 'Task: Help me complete Android Firebase registration.')
add_code(doc, '')
add_code(doc, '1. Show me the exact flutterfire CLI command to run locally to register')
add_code(doc, '   the Android app (bundle ID: com.ziegler.logbook or confirm from')
add_code(doc, '   android/app/build.gradle).')
add_code(doc, '2. Once I have run flutterfire and have google-services.json:')
add_code(doc, '   - Confirm where to place google-services.json')
add_code(doc, '   - Update firebase_options.dart with the real Android appId')
add_code(doc, '   - Verify android/app/build.gradle has the correct applicationId and')
add_code(doc, '     the google-services plugin applied')
add_code(doc, '3. Verify the android/ directory has the necessary plugin classpath in')
add_code(doc, '   android/build.gradle or settings.gradle.kts')
add_code(doc, '')
add_code(doc, 'Commit and push changes to main once verified.')

doc.add_paragraph()
add_divider(doc)

# A3
add_heading(doc, 'A3 — Multilingual (German / English)', level=2)
add_para(doc, 'Effort: 11 days | Dependency: None | Emergency Manifest stays English.', italic=True, color=SLATE)
doc.add_paragraph()
add_para(doc, 'Paste this prompt into a new session:', bold=True)
add_code(doc, 'We are working on the Flutter sailing logbook app at /home/user/Logbook.')
add_code(doc, 'The app is currently hardcoded to German (de_CH). We want to support')
add_code(doc, 'German and English, with the Emergency Manifest screen permanently')
add_code(doc, 'in English (do not extract strings from emergency_manifest_screen.dart).')
add_code(doc, '')
add_code(doc, 'flutter_localizations and intl are already in pubspec.yaml.')
add_code(doc, '')
add_code(doc, 'CRITICAL — fix the Besatzung sentinel first:')
add_code(doc, 'The string "Besatzung: " in home_repository.dart is used both as display')
add_code(doc, 'text and as a detection sentinel (startsWith). Before extracting strings,')
add_code(doc, 'change the stored prefix to a non-localised ASCII key "crew:" and create')
add_code(doc, 'a separate display string for the ARB file. Update the detection logic in')
add_code(doc, 'day_detail_screen.dart and home_repository.dart accordingly.')
add_code(doc, '')
add_code(doc, 'Then:')
add_code(doc, '1. Add l10n.yaml (arb-dir: lib/l10n, template-arb-file: app_de.arb)')
add_code(doc, '2. Create lib/l10n/app_de.arb with all extracted German strings')
add_code(doc, '3. Create lib/l10n/app_en.arb with English translations')
add_code(doc, '4. Wire AppLocalizations.delegate in lib/app.dart')
add_code(doc, '5. Add language selector in Settings screen, persist choice in ThemeProvider')
add_code(doc, '6. Verify DateFormat calls use the active locale')
add_code(doc, '')
add_code(doc, 'Skip all files in lib/features/emergency/.')
add_code(doc, 'Commit and push to main after each phase.')

doc.add_paragraph()
add_divider(doc)

# A4
add_heading(doc, 'A4 — Authentication (Email, Apple, Google)', level=2)
add_para(doc, 'Effort: 20 days | Dependency: Android Firebase registered (A2) | New Firestore structure.', italic=True, color=SLATE)
doc.add_paragraph()
add_para(doc, 'Paste this prompt into a new session:', bold=True)
add_code(doc, 'We are working on the Flutter sailing logbook app at /home/user/Logbook.')
add_code(doc, 'firebase_auth is in pubspec.yaml but never wired up. We want to add:')
add_code(doc, '  - Email/Password login')
add_code(doc, '  - Apple Sign-In (iOS + macOS)')
add_code(doc, '  - Google Sign-In (iOS + Android)')
add_code(doc, '')
add_code(doc, 'Current Firestore path: logbooks/{installationId}/...')
add_code(doc, 'Target Firestore path:  boats/{boatId}/...  +  users/{uid}/profile')
add_code(doc, '')
add_code(doc, 'There are NO existing users to migrate. Clean cutover.')
add_code(doc, '')
add_code(doc, 'Phase 1 — Platform setup (tell me what to do in Firebase Console and Xcode):')
add_code(doc, '  - Enable Email/Password, Apple, Google in Firebase Auth Console')
add_code(doc, '  - sign_in_with_apple capability for iOS and macOS targets')
add_code(doc, '  - Google reversed-client-id URL scheme in Info.plist')
add_code(doc, '  - SHA-1 fingerprint for Android in Firebase Console')
add_code(doc, '  - Add packages: sign_in_with_apple, google_sign_in to pubspec.yaml')
add_code(doc, '')
add_code(doc, 'Phase 2 — Auth UI:')
add_code(doc, '  - Login screen: email+password fields, Apple button, Google button')
add_code(doc, '  - Register screen, forgot-password screen')
add_code(doc, '  - Auth guard in GoRouter (router.dart) — redirect unauthenticated to /login')
add_code(doc, '  - Account section in settings_screen.dart: email, sign out, delete account')
add_code(doc, '')
add_code(doc, 'Phase 3 — Data model:')
add_code(doc, '  - Update lib/core/services/firestore_service.dart to use boatId')
add_code(doc, '    (from users/{uid}/profile.boatId) instead of installationId')
add_code(doc, '  - On first authenticated launch: create users/{uid} and boats/{boatId}')
add_code(doc, '  - Write Firestore Security Rules: only members of boats/{boatId} can r/w')
add_code(doc, '')
add_code(doc, 'Phase 4 — Multi-device share:')
add_code(doc, '  - "Connect to existing logbook" flow: user enters boatId')
add_code(doc, '  - Cloud Function or client-side: add uid to boats/{boatId}/members')
add_code(doc, '')
add_code(doc, 'Commit and push each phase to main separately.')

doc.add_paragraph()
add_divider(doc)

# A5
add_heading(doc, 'A5 — Web Subscription + Token Billing (optional, implement later)', level=2)
add_para(doc, 'Effort: 2 weeks | Dependency: Auth (A4) must be complete | Only implement if Firebase costs become meaningful.', italic=True, color=SLATE)
doc.add_paragraph()
add_para(doc, 'Paste this prompt into a new session:', bold=True)
add_code(doc, 'We are working on the Flutter sailing logbook app at /home/user/Logbook.')
add_code(doc, 'Auth is complete (Firebase Auth, users/{uid} in Firestore).')
add_code(doc, 'We want to add optional paid cloud sync using a web-based subscription')
add_code(doc, 'model to avoid App Store IAP (no 30% platform cut, no IAP review).')
add_code(doc, '')
add_code(doc, 'Model:')
add_code(doc, '  1. User pays on website (Stripe) → webhook → Cloud Function generates UUID')
add_code(doc, '     token → emails it to user')
add_code(doc, '  2. User enters token in app Settings ("Sync-Code aktivieren")')
add_code(doc, '  3. App sends token to validation Cloud Function → sets')
add_code(doc, '     users/{uid}/subscription: {status: "active", expiresAt}')
add_code(doc, '  4. Firestore Security Rules gate write access on subscription.status')
add_code(doc, '')
add_code(doc, 'Tasks:')
add_code(doc, '  - Write Firebase Cloud Functions: generateToken (Stripe webhook),')
add_code(doc, '    validateToken (called from app)')
add_code(doc, '  - Add token entry field + "Activate" button to settings_screen.dart')
add_code(doc, '  - Add subscription status display in settings_screen.dart')
add_code(doc, '  - Add sync gate in home_repository.dart: check subscription before')
add_code(doc, '    any Firestore write; fall back to local-only if not subscribed')
add_code(doc, '  - Add persistent banner in home_screen.dart when sync is inactive')
add_code(doc, '  - Update Firestore Security Rules to check subscription.status')
add_code(doc, '')
add_code(doc, 'Commit and push each component to main separately.')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX B — FIREBASE COST REFERENCE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Appendix B — Firebase Cost Reference', level=1)

add_heading(doc, 'Spark plan (free) limits', level=2)
make_table(doc,
    ['Service', 'Free limit', 'What exceeds it'],
    [
        ['Firestore reads', '50,000 / day', 'Active users doing many reads on app open'],
        ['Firestore writes', '20,000 / day', 'Heavy logging days with many timeline entries'],
        ['Firestore storage', '1 GiB total', 'Approximately 1,300 users with a full year of entries'],
        ['Firebase Storage', '5 GiB total', 'Approximately 500 users with 10 MB of GPX tracks each'],
        ['Firebase Auth (email)', '10,000 MAU / month', 'Not a concern at early scale'],
        ['Auth (Google / Apple)', 'Unlimited on Identity Platform', 'Always free'],
        ['Cloud Functions invocations', '125,000 / month', 'Only relevant if billing webhooks are added'],
    ],
    col_widths=[1.7, 1.6, 3.3],
)

doc.add_paragraph()
add_heading(doc, 'Blaze pricing (after free tier)', level=2)
make_table(doc,
    ['Service', 'Price'],
    [
        ['Firestore reads', '$0.06 per 100,000'],
        ['Firestore writes', '$0.18 per 100,000'],
        ['Firestore storage', '$0.18 per GiB/month'],
        ['Firebase Storage', '$0.026 per GiB/month'],
        ['Cloud Functions', '$0.40 per million invocations + compute time'],
    ],
    col_widths=[2.5, 4.1],
)

doc.add_paragraph()
add_callout(doc, 'Budget protection:',
    'In Firebase Console → Project Settings → Usage and billing → Set a budget alert at €20/month '
    'and a spending limit at €50/month. The limit prevents runaway costs but will cause Firestore '
    'writes to fail once hit — users will see sync errors rather than unexpected charges.',
    color=AMBER)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX C — DECISIONS LOG
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Appendix C — Key Architectural Decisions Log', level=1)
add_para(doc, 'This log records decisions made during planning so they do not need to be re-litigated in future sessions.')

doc.add_paragraph()
make_table(doc,
    ['Decision', 'Choice made', 'Reasoning'],
    [
        ['Emergency Manifest language',
         'Always English',
         'Already written in English; maritime safety document used internationally; '
         'do not localise'],
        ['Auth providers',
         'Email/Password + Apple + Google',
         'Apple Sign-In mandatory when offering any social login on iOS (App Store rule). '
         'Google Sign-In covers Android. Email/Password as universal fallback.'],
        ['Billing model (current)',
         'No billing — absorb Firebase costs',
         'Firebase cost is < €5/month at 1,000 users. Billing complexity is not justified '
         'at current scale. Revisit if cost exceeds €20/month.'],
        ['Billing model (if needed)',
         'Web subscription + token (not App Store IAP)',
         'Avoids 30% platform cut, no App Store billing review, acceptable under App Store '
         'guidelines when sync is framed as an external service.'],
        ['Map tile provider',
         'MapTiler (nautical style) + Esri (satellite)',
         'OSM demo tile server is policy-prohibited for published apps. '
         'MapTiler free tier covers early scale; nautical chart style is relevant for sailors.'],
        ['State management',
         'Keep provider package — no migration',
         'Existing ChangeNotifier providers are correct for the app\'s complexity level. '
         'Auth and subscription each add one new provider to MultiProvider.'],
        ['Data model',
         'logbooks/{id} → boats/{boatId} + users/{uid}',
         'Enables Firestore Security Rules based on auth. No migration needed (no existing users).'],
        ['Multi-device share',
         'Keep code-based sharing via boatId',
         'Preserves the elegant "share a boat" UX. Second device enters boatId after login; '
         'uid is added to boats/{boatId}/members.'],
        ['Android state',
         'Firebase not yet configured (placeholder)',
         'Must run flutterfire configure --platforms=android and add google-services.json '
         'before Android testing. This is Phase 2.1 of the auth track.'],
        ['Grace period for old users',
         'None',
         'Developer is the only current user. Clean cutover to new data model.'],
        ['i18n sentinel fix',
         '"crew:" prefix in logic, localised display string in ARB',
         '"Besatzung: " is used as both display text and detection marker. Splitting these '
         'prevents localisation from breaking crew-change detection.'],
    ],
    col_widths=[1.6, 1.5, 3.5],
)

doc.add_paragraph()
doc.add_paragraph()
add_para(doc,
    f'Document generated: {datetime.datetime.now().strftime("%d %B %Y, %H:%M")}  ·  '
    'App repository: cpahab/Logbook  ·  Current branch: main',
    italic=True,
    color=RGBColor(0x99, 0xAA, 0xBB),
    size=9,
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/user/Logbook/docs/Logbook_Upgrade_Plan.docx'
doc.save(out)
print(f'Saved → {out}')
